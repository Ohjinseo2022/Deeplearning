from openpyxl import load_workbook  # 엑셀파일을 불러오기 위한 라이브러리
import docx
from docx.oxml.ns import qn  # 한글 폰트를 인식시키기 위한 라이브러리
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 텍스트 정렬을 위한 라이브러리
from docx2pdf import convert  # Docx 파일을 pdf 파일로 변환하기 위한 라이브러리

# load_workbook 에서 사용할수 있는 속성값
# -filename
# -read_only  읽기만 해서 쓰기제한을 둘것인지 설정하는 파라미터입니다. 아무 언급을 안하면 쓰기도 가능한 상태로 엑셀을 불러들입니다.
# -keep_vba 엑셀에 있는 vba(액셀의 매크로기능 Visual Basic for Applications) 기능을 살려서 가져올것인지 아닌지를 설정합니다. 기본적으로 엑셀에 저장되어 있는 vba은 가져오지 않는 것으로 설정되어있습니다
# -data_only  엑셀의 수식값을 가져올것인지 수식의 결과인 data를 가져올지를 묻는 설정입니다.
# -keep_links 엑셀파일안에 있는 링크를 그대로 쓸것인지를 설정할 수 있습니다.
load_wb = load_workbook(r"수료증명단.xlsx")
load_ws = load_wb.active

name_list = []
birthday_list = []
ho_list = []

## 각각의 리스트에 이름, 생일, 번호 순서대로 접근하여 추가
for i in range(1, load_ws.max_row + 1):
    name_list.append(load_ws.cell(i, 1).value)
    birthday_list.append(load_ws.cell(i, 2).value)
    ho_list.append(load_ws.cell(i, 3).value)

## 제대로 데이터가 들어갔는지 확인용
print(name_list)
print(birthday_list)
print(ho_list)

for i in range(len(name_list)):
    doc = docx.Document(r"수료증양식.docx")

    style = doc.styles["Normal"]
    style.font.name = "나눔고딕"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
    style.font.size = docx.shared.Pt(12)

    para = doc.add_paragraph()
    run = para.add_run(f"제 {ho_list[i]}호")
    run.font.name = "나눔고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run("\n\n")
    run = para.add_run("수  료  증")
    run.font.name = "나눔고딕"
    run.bold = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
    run.font.size = docx.shared.Pt(40)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 가운데 정렬

    para = doc.add_paragraph()
    run = para.add_run("\n\n")
    run = para.add_run(f"성 명 : {name_list[i]} \n")
    run.font.name = "나눔고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(f"생 년 월 일 : {birthday_list[i]} \n")
    run.font.name = "나눔고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
    run.font.size = docx.shared.Pt(20)
    run = para.add_run("교 육 과 정 : 파이썬과 40개의 작품들\n")
    run.font.name = "나눔고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
    run.font.size = docx.shared.Pt(20)
    run = para.add_run("교 육 날 짜 : 2021.08.05~2021.09.09\n")
    run.font.name = "나눔고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run("\n\n")
    run = para.add_run(" 위 사람은 파이썬과 40개의 작품들 교육과정을 \n")
    run.font.name = "나눔고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(" 이수하였으므로 이 증서를 수여 합니다 \n")
    run.font.name = "나눔고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run("\n\n")
    run = para.add_run("파이썬교육기관장\n")
    run.font.name = "나눔고딕"
    run.bold = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "나눔고딕")
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 가운데 정렬

    doc.save("{}.docx".format(name_list[i]))
    convert("{}.docx".format(name_list[i]), "./{}.pdf".format(name_list[i]))
